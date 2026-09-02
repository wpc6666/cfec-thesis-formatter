#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
重庆财经学院本科毕业论文 DOCX 生成引擎
依据：《重庆财经学院毕业论文（设计）指导手册》
依赖：python-docx >= 0.8.11
用法：
    from cfec_thesis_docx import ThesisFormatter
    formatter = ThesisFormatter()
    formatter.build(thesis_data)
    formatter.save("output.docx")
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ============================================================
#  格式常量（全部来自重庆财经学院指导手册）
# ============================================================
class FormatSpec:
    # 页面
    PAPER_WIDTH = Cm(21)
    PAPER_HEIGHT = Cm(29.7)
    MARGIN_LEFT = Cm(2.8)
    MARGIN_RIGHT = Cm(2.5)
    MARGIN_TOP = Cm(2.5)
    MARGIN_BOTTOM = Cm(2.5)
    HEADER_DISTANCE = Cm(1.5)
    FOOTER_DISTANCE = Cm(1.75)

    # 行距：固定值26磅
    LINE_SPACING_PT = 26

    # 字体
    FONT_SONG = "宋体"
    FONT_HEI = "黑体"
    FONT_TNR = "Times New Roman"

    # 字号映射（中文字号 -> 磅值）
    SIZE_XIAO_SAN = Pt(15)    # 小三
    SIZE_SI = Pt(14)           # 四号
    SIZE_XIAO_SI = Pt(12)     # 小四
    SIZE_WU = Pt(10.5)         # 五号
    SIZE_XIAO_WU = Pt(9)       # 小五

    # 标题
    H1_SIZE = SIZE_XIAO_SAN
    H2_SIZE = SIZE_SI
    H3_SIZE = SIZE_XIAO_SI
    H4_SIZE = SIZE_XIAO_SI

    # 缩进
    INDENT_2CHAR = Cm(0.74) * 2  # 约2字符（小四12pt，1字符≈0.42cm）


# ============================================================
#  工具函数
# ============================================================
def set_run_font(run, cn_font=FormatSpec.FONT_SONG, en_font=FormatSpec.FONT_TNR,
                 size=FormatSpec.SIZE_XIAO_SI, bold=False, color=None):
    """设置run的中英文字体、字号、加粗、颜色"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = en_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    if color:
        run.font.color.rgb = color


def set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=None, line_spacing_pt=FormatSpec.LINE_SPACING_PT,
                         space_before=0, space_after=0, left_indent=None):
    """设置段落格式：对齐、首行缩进、固定行距、段前段后"""
    pf = para.paragraph_format
    pf.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def add_field(paragraph, field_code):
    """在段落中插入Word域（用于目录、页码等）"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)
    return run


def set_cell_border(cell, **kwargs):
    """设置单元格边框（用于三线表）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ['sz', 'val', 'color', 'space']:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_section_page_number(section, fmt='decimal', start=1):
    """设置节的页码格式和起始编号
    fmt: 'decimal' 阿拉伯数字, 'upperRoman' 大写罗马数字
    """
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))


def unlink_header_footer(section):
    """取消节的页眉页脚'链接到前一条'"""
    for header in section.header.part.related_parts.values():
        pass
    # 通过XML设置
    sectPr = section._sectPr
    # 页眉引用
    headerReferences = sectPr.findall(qn('w:headerReference'))
    for href in headerReferences:
        pass
    # 设置titlePg不继承
    titlePg = sectPr.find(qn('w:titlePg'))
    if titlePg is None:
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)


# ============================================================
#  论文格式化器主类
# ============================================================
class ThesisFormatter:
    """重庆财经学院本科毕业论文 DOCX 生成器"""

    def __init__(self):
        self.doc = Document()
        self._setup_default_style()
        self._setup_page()
        self.current_chapter = 0
        self.figure_counter = 0
        self.table_counter = 0
        self.equation_counter = 0

    def _setup_default_style(self):
        """设置文档默认样式"""
        style = self.doc.styles['Normal']
        style.font.name = FormatSpec.FONT_TNR
        style.font.size = FormatSpec.SIZE_XIAO_SI
        style.element.rPr.rFonts.set(qn('w:eastAsia'), FormatSpec.FONT_SONG)
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(FormatSpec.LINE_SPACING_PT)

    def _setup_page(self):
        """页面设置：A4，页边距，页眉页脚距离"""
        section = self.doc.sections[0]
        section.page_width = FormatSpec.PAPER_WIDTH
        section.page_height = FormatSpec.PAPER_HEIGHT
        section.left_margin = FormatSpec.MARGIN_LEFT
        section.right_margin = FormatSpec.MARGIN_RIGHT
        section.top_margin = FormatSpec.MARGIN_TOP
        section.bottom_margin = FormatSpec.MARGIN_BOTTOM
        section.header_distance = FormatSpec.HEADER_DISTANCE
        section.footer_distance = FormatSpec.FOOTER_DISTANCE

    # ----------------------------------------------------------
    #  封面
    # ----------------------------------------------------------
    def create_cover(self, title, college="", grade="", major="",
                     student_id="", name="", advisor="", date="二〇二六年六月"):
        """生成封面页（无页码）"""
        # 封面信息居中
        for _ in range(3):
            self.doc.add_paragraph()

        # 学校名
        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=30)
        run = p.add_run("重庆财经学院")
        set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_SI, bold=True)

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=30)
        run = p.add_run("本科毕业论文（设计）")
        set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_SI, bold=True)

        for _ in range(2):
            self.doc.add_paragraph()

        # 论文题目
        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=30)
        run = p.add_run("论文题目：")
        set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_SI, bold=True)

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=30)
        run = p.add_run(title)
        set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_SI, bold=True)

        for _ in range(2):
            self.doc.add_paragraph()

        # 信息栏
        info_items = [
            ("学　　院：", college),
            ("年　　级：", grade),
            ("专　　业：", major),
            ("学　　号：", student_id),
            ("姓　　名：", name),
            ("指导教师：", advisor),
        ]
        for label, value in info_items:
            p = self.doc.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=28)
            run = p.add_run(f"{label}{value if value else '　' * 8}")
            set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_SI, bold=True)

        for _ in range(2):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=30)
        run = p.add_run(date)
        set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_SI, bold=True)

        # 封面无页码：设置本节页码格式但不显示
        section = self.doc.sections[0]
        set_section_page_number(section, fmt='decimal', start=0)
        # 封面页脚清空
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()

    # ----------------------------------------------------------
    #  分节：摘要/目录部分（罗马数字页码）
    # ----------------------------------------------------------
    def _start_front_matter_section(self):
        """开始前置部分（摘要、目录），大写罗马数字页码"""
        new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.page_width = FormatSpec.PAPER_WIDTH
        new_section.page_height = FormatSpec.PAPER_HEIGHT
        new_section.left_margin = FormatSpec.MARGIN_LEFT
        new_section.right_margin = FormatSpec.MARGIN_RIGHT
        new_section.top_margin = FormatSpec.MARGIN_TOP
        new_section.bottom_margin = FormatSpec.MARGIN_BOTTOM
        new_section.header_distance = FormatSpec.HEADER_DISTANCE
        new_section.footer_distance = FormatSpec.FOOTER_DISTANCE
        new_section.footer.is_linked_to_previous = False
        set_section_page_number(new_section, fmt='upperRoman', start=1)
        self._add_page_number_to_footer(new_section)
        return new_section

    def _start_body_section(self):
        """开始正文部分，阿拉伯数字页码从1开始"""
        new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.page_width = FormatSpec.PAPER_WIDTH
        new_section.page_height = FormatSpec.PAPER_HEIGHT
        new_section.left_margin = FormatSpec.MARGIN_LEFT
        new_section.right_margin = FormatSpec.MARGIN_RIGHT
        new_section.top_margin = FormatSpec.MARGIN_TOP
        new_section.bottom_margin = FormatSpec.MARGIN_BOTTOM
        new_section.header_distance = FormatSpec.HEADER_DISTANCE
        new_section.footer_distance = FormatSpec.FOOTER_DISTANCE
        new_section.footer.is_linked_to_previous = False
        set_section_page_number(new_section, fmt='decimal', start=1)
        self._add_page_number_to_footer(new_section)
        return new_section

    def _add_page_number_to_footer(self, section):
        """在页脚居中添加页码域"""
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run()
        set_run_font(run, en_font=FormatSpec.FONT_TNR, size=FormatSpec.SIZE_XIAO_SI)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

    # ----------------------------------------------------------
    #  中文摘要
    # ----------------------------------------------------------
    def create_abstract_cn(self, content, keywords):
        """中文摘要页"""
        # 标题
        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
        run = p.add_run("摘　要")  # 两字间空2个半角空格
        set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_XIAO_SAN, bold=True)

        # 内容
        p = self.doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=FormatSpec.INDENT_2CHAR)
        run = p.add_run(content)
        set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_XIAO_SI)

        # 空一行
        self.doc.add_paragraph()

        # 关键词
        p = self.doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=FormatSpec.INDENT_2CHAR)
        run = p.add_run("关键词：")
        set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_XIAO_SI)
        run = p.add_run("；".join(keywords))
        set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_XIAO_SI)

    # ----------------------------------------------------------
    #  英文摘要
    # ----------------------------------------------------------
    def create_abstract_en(self, content, keywords):
        """英文摘要页（Abstract）"""
        self.doc.add_page_break()

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
        run = p.add_run("Abstract")
        set_run_font(run, en_font=FormatSpec.FONT_TNR, size=FormatSpec.SIZE_XIAO_SAN, bold=True)

        p = self.doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=FormatSpec.INDENT_2CHAR)
        run = p.add_run(content)
        set_run_font(run, en_font=FormatSpec.FONT_TNR, size=FormatSpec.SIZE_XIAO_SI)

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=FormatSpec.INDENT_2CHAR)
        run = p.add_run("Keywords: ")
        set_run_font(run, en_font=FormatSpec.FONT_TNR, size=FormatSpec.SIZE_XIAO_SI, bold=True)
        run = p.add_run("; ".join(keywords))
        set_run_font(run, en_font=FormatSpec.FONT_TNR, size=FormatSpec.SIZE_XIAO_SI)

    # ----------------------------------------------------------
    #  目录
    # ----------------------------------------------------------
    def create_toc(self):
        """目录页（只显示到二级标题，点号前导符）"""
        self.doc.add_page_break()

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
        run = p.add_run("目　录")
        set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_XIAO_SAN, bold=True)

        # 插入TOC域，显示级别2
        p = self.doc.add_paragraph()
        set_paragraph_format(p, line_spacing_pt=FormatSpec.LINE_SPACING_PT)
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = r'TOC \o "1-2" \h \z \u'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        # 占位文本
        placeholder = OxmlElement('w:r')
        placeholder_t = OxmlElement('w:t')
        placeholder_t.text = "（请在Word中按F9更新目录）"
        placeholder.append(placeholder_t)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        run._element.append(placeholder)
        run._element.append(fldChar3)

    # ----------------------------------------------------------
    #  四级标题
    # ----------------------------------------------------------
    def add_h1(self, text):
        """一级标题：一、 黑体小三加粗居中，与上文空一行"""
        self.current_chapter += 1
        self.figure_counter = 0
        self.table_counter = 0
        self.equation_counter = 0

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_before=12, space_after=6)
        # 应用Heading 1样式（用于目录）
        p.style = self.doc.styles['Heading 1']
        run = p.add_run(f"{self._chinese_num(self.current_chapter)}、{text}")
        set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_XIAO_SAN, bold=True,
                     color=RGBColor(0, 0, 0))
        return p

    def add_h2(self, text):
        """二级标题：（一） 黑体四号首行缩进2字符"""
        cnt = self._get_counter('h2')
        p = self.doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=FormatSpec.INDENT_2CHAR,
                             space_before=6, space_after=3)
        p.style = self.doc.styles['Heading 2']
        run = p.add_run(f"（{self._chinese_num(cnt)}）{text}")
        set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_SI,
                     color=RGBColor(0, 0, 0))
        return p

    def add_h3(self, text):
        """三级标题：1. 黑体小四号首行缩进2字符"""
        cnt = self._get_counter('h3')
        p = self.doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=FormatSpec.INDENT_2CHAR,
                             space_before=3, space_after=3)
        run = p.add_run(f"{cnt}. {text}")
        set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_XIAO_SI)
        return p

    def add_h4(self, text):
        """四级标题：（1） 宋体小四号加粗首行缩进2字符"""
        cnt = self._get_counter('h4')
        p = self.doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=FormatSpec.INDENT_2CHAR,
                             space_before=3, space_after=3)
        run = p.add_run(f"（{cnt}）{text}")
        set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_XIAO_SI, bold=True)
        return p

    # ----------------------------------------------------------
    #  正文段落
    # ----------------------------------------------------------
    def add_body(self, text):
        """正文：宋体小四，两端对齐，首行缩进2字符，固定26磅"""
        p = self.doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=FormatSpec.INDENT_2CHAR)
        run = p.add_run(text)
        set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_XIAO_SI)
        return p

    def add_body_with_citations(self, text, citations=None):
        """正文带参考文献引用上标
        citations: list of (start_pos, end_pos, ref_text) 或简单的引用标记
        """
        p = self.doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=FormatSpec.INDENT_2CHAR)
        # 简单实现：将 [n] 标记转为上标
        import re
        parts = re.split(r'(\[\d+(?:[-,]\d+)*\])', text)
        for part in parts:
            if re.match(r'\[\d+(?:[-,]\d+)*\]', part):
                run = p.add_run(part)
                set_run_font(run, en_font=FormatSpec.FONT_TNR, size=FormatSpec.SIZE_XIAO_SI)
                run.font.superscript = True
            else:
                run = p.add_run(part)
                set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_XIAO_SI)
        return p

    # ----------------------------------------------------------
    #  图
    # ----------------------------------------------------------
    def add_figure(self, image_path, caption, width_cm=12):
        """图：图题在图下方，黑体五号居中，编号 图 1-1"""
        self.figure_counter += 1
        fig_num = f"图 {self.current_chapter}-{self.figure_counter}"

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
        if os.path.exists(image_path):
            run = p.add_run()
            run.add_picture(image_path, width=Cm(width_cm))
        else:
            run = p.add_run(f"[图片占位：{image_path}]")
            set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_WU,
                         color=RGBColor(128, 128, 128))

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        run = p.add_run(f"{fig_num} {caption}")
        set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_WU)

    # ----------------------------------------------------------
    #  三线表
    # ----------------------------------------------------------
    def add_three_line_table(self, headers, rows, caption):
        """三线表：表题在表上方，黑体五号居中，编号 表 1-1"""
        self.table_counter += 1
        table_num = f"表 {self.current_chapter}-{self.table_counter}"

        # 表题
        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6)
        run = p.add_run(f"{table_num} {caption}")
        set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_WU)

        # 创建表格
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # 表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = ""
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_WU, bold=True)
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 数据行
        for row_idx, row_data in enumerate(rows):
            cells = table.rows[row_idx + 1].cells
            for col_idx, cell_text in enumerate(row_data):
                cells[col_idx].text = ""
                p = cells[col_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(cell_text))
                set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_WU)
                cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 三线表边框：顶线、表头底线、底线
        for row in table.rows:
            for cell in row.cells:
                # 先清除所有边框
                set_cell_border(cell,
                    top={"val": "nil"}, bottom={"val": "nil"},
                    left={"val": "nil"}, right={"val": "nil"})

        # 顶线（第一行顶部）
        for cell in table.rows[0].cells:
            set_cell_border(cell, top={"val": "single", "sz": "12", "color": "000000"})
        # 表头底线（第一行底部）
        for cell in table.rows[0].cells:
            set_cell_border(cell, bottom={"val": "single", "sz": "6", "color": "000000"})
        # 底线（最后一行底部）
        for cell in table.rows[-1].cells:
            set_cell_border(cell, bottom={"val": "single", "sz": "12", "color": "000000"})

        # 表后空行
        p = self.doc.add_paragraph()
        set_paragraph_format(p, space_after=6)

    # ----------------------------------------------------------
    #  公式
    # ----------------------------------------------------------
    def add_equation(self, equation_text, caption=None):
        """公式：单独成行，TNR小四，编号（1.1）右对齐"""
        self.equation_counter += 1
        eq_num = f"（{self.current_chapter}.{self.equation_counter}）"

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=6)
        # 公式内容
        run = p.add_run(equation_text + "　　　　")
        set_run_font(run, en_font=FormatSpec.FONT_TNR, size=FormatSpec.SIZE_XIAO_SI)
        # 编号右对齐（用制表符）
        run = p.add_run(eq_num)
        set_run_font(run, en_font=FormatSpec.FONT_TNR, size=FormatSpec.SIZE_XIAO_SI)

    # ----------------------------------------------------------
    #  参考文献
    # ----------------------------------------------------------
    def add_references(self, references):
        """参考文献：另起一页，悬挂缩进2字符，五号字，编号[1]"""
        self.doc.add_page_break()

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
        run = p.add_run("参考文献")
        set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_XIAO_SAN, bold=True)

        for idx, ref in enumerate(references, 1):
            p = self.doc.add_paragraph()
            # 悬挂缩进：左缩进2字符，首行缩进-2字符
            pf = p.paragraph_format
            pf.left_indent = FormatSpec.INDENT_2CHAR
            pf.first_line_indent = Cm(-0.74) * 2
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(FormatSpec.LINE_SPACING_PT)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

            run = p.add_run(f"[{idx}] {ref}")
            set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_WU)

    # ----------------------------------------------------------
    #  附录
    # ----------------------------------------------------------
    def add_appendix(self, label, title, content):
        """附录：附录A 黑体小三加粗居中，标题黑体四号居中"""
        self.doc.add_page_break()

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6)
        run = p.add_run(f"附录　{label}")
        set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_XIAO_SAN, bold=True)

        p = self.doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
        run = p.add_run(title)
        set_run_font(run, cn_font=FormatSpec.FONT_HEI, size=FormatSpec.SIZE_SI)

        p = self.doc.add_paragraph()
        set_paragraph_format(p, first_line_indent=FormatSpec.INDENT_2CHAR)
        run = p.add_run(content)
        set_run_font(run, cn_font=FormatSpec.FONT_SONG, size=FormatSpec.SIZE_XIAO_SI)

    # ----------------------------------------------------------
    #  完整构建
    # ----------------------------------------------------------
    def build(self, thesis_data):
        """根据论文数据字典完整构建文档
        thesis_data 结构：
        {
            "cover": {"title": "...", "college": "...", ...},
            "abstract_cn": {"content": "...", "keywords": [...]},
            "abstract_en": {"content": "...", "keywords": [...]},
            "chapters": [
                {"title": "绪论", "sections": [
                    {"title": "研究背景", "subsections": [
                        {"title": "...", "paragraphs": ["...", "..."]}
                    ], "paragraphs": ["..."]}
                ]}
            ],
            "references": ["...", "..."],
            "appendices": [{"label": "A", "title": "...", "content": "..."}]
        }
        """
        # 1. 封面
        cover = thesis_data.get("cover", {})
        self.create_cover(
            title=cover.get("title", "论文题目"),
            college=cover.get("college", ""),
            grade=cover.get("grade", ""),
            major=cover.get("major", ""),
            student_id=cover.get("student_id", ""),
            name=cover.get("name", ""),
            advisor=cover.get("advisor", ""),
            date=cover.get("date", "二〇二六年六月")
        )

        # 2. 前置部分分节（罗马数字页码）
        self._start_front_matter_section()

        # 3. 中文摘要
        abs_cn = thesis_data.get("abstract_cn", {})
        self.create_abstract_cn(
            content=abs_cn.get("content", "（中文摘要内容）"),
            keywords=abs_cn.get("keywords", ["关键词1", "关键词2"])
        )

        # 4. 英文摘要
        abs_en = thesis_data.get("abstract_en", {})
        self.create_abstract_en(
            content=abs_en.get("content", "(English abstract content)"),
            keywords=abs_en.get("keywords", ["keyword1", "keyword2"])
        )

        # 5. 目录
        self.create_toc()

        # 6. 正文分节（阿拉伯数字页码）
        self._start_body_section()

        # 7. 各章节
        chapters = thesis_data.get("chapters", [])
        for chapter in chapters:
            self.add_h1(chapter.get("title", "章节标题"))
            for section in chapter.get("sections", []):
                self.add_h2(section.get("title", "二级标题"))
                for para in section.get("paragraphs", []):
                    if "[" in para and "]" in para:
                        self.add_body_with_citations(para)
                    else:
                        self.add_body(para)
                for subsection in section.get("subsections", []):
                    self.add_h3(subsection.get("title", "三级标题"))
                    for para in subsection.get("paragraphs", []):
                        if "[" in para and "]" in para:
                            self.add_body_with_citations(para)
                        else:
                            self.add_body(para)

        # 8. 参考文献
        refs = thesis_data.get("references", [])
        if refs:
            self.add_references(refs)

        # 9. 附录
        appendices = thesis_data.get("appendices", [])
        for app in appendices:
            self.add_appendix(
                label=app.get("label", "A"),
                title=app.get("title", "附录标题"),
                content=app.get("content", "附录内容")
            )

        return self.doc

    # ----------------------------------------------------------
    #  工具方法
    # ----------------------------------------------------------
    def _chinese_num(self, n):
        """阿拉伯数字转中文数字（1-99）"""
        digits = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九"]
        if n < 10:
            return digits[n]
        elif n < 20:
            return "十" + (digits[n % 10] if n % 10 != 0 else "")
        elif n < 100:
            tens = n // 10
            ones = n % 10
            return digits[tens] + "十" + (digits[ones] if ones != 0 else "")
        return str(n)

    def _get_counter(self, level):
        """获取标题计数器（简单实现，基于文档中已有标题数）"""
        if level == 'h2':
            if not hasattr(self, '_h2_counter'):
                self._h2_counter = 0
            self._h2_counter += 1
            return self._h2_counter
        elif level == 'h3':
            if not hasattr(self, '_h3_counter'):
                self._h3_counter = 0
            self._h3_counter += 1
            return self._h3_counter
        elif level == 'h4':
            if not hasattr(self, '_h4_counter'):
                self._h4_counter = 0
            self._h4_counter += 1
            return self._h4_counter
        return 1

    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
        return output_path


# ============================================================
#  命令行使用示例
# ============================================================
if __name__ == "__main__":
    import json

    # 示例论文数据
    sample_thesis = {
        "cover": {
            "title": "基于Python的毕业论文格式自动化生成系统研究",
            "college": "会计学院",
            "grade": "2022级",
            "major": "会计学",
            "student_id": "2022010101",
            "name": "张三",
            "advisor": "李四教授",
            "date": "二〇二六年六月"
        },
        "abstract_cn": {
            "content": "本文针对重庆财经学院本科毕业论文格式规范要求，设计并实现了一套基于python-docx的自动化排版系统。该系统能够将结构化论文数据一键转换为符合学校规范的Word文档，涵盖封面、摘要、目录、正文、图表、公式、参考文献和附录等全部要素。实验表明，该系统能够显著减少人工排版时间，格式准确率达到98%以上。",
            "keywords": ["毕业论文", "格式规范", "python-docx", "自动化排版"]
        },
        "abstract_en": {
            "content": "This paper designs and implements an automatic typesetting system based on python-docx for the undergraduate thesis format requirements of Chongqing Finance and Economics College. The system can convert structured thesis data into Word documents conforming to school specifications with one click, covering all elements such as cover, abstract, table of contents, body, figures, tables, formulas, references and appendices.",
            "keywords": ["thesis", "format specification", "python-docx", "automatic typesetting"]
        },
        "chapters": [
            {
                "title": "绪论",
                "sections": [
                    {
                        "title": "研究背景",
                        "paragraphs": [
                            "毕业论文是本科教学的重要环节，其格式规范性直接影响论文质量。重庆财经学院对毕业论文格式有严格要求[1]，包括页边距、行距、字体字号、页码规则、标题层级等多个方面。",
                            "传统人工排版方式耗时费力，且容易出现格式不一致的问题[2-3]。因此，开发一套自动化排版系统具有重要的现实意义。"
                        ]
                    },
                    {
                        "title": "研究意义",
                        "subsections": [
                            {
                                "title": "理论意义",
                                "paragraphs": ["本文丰富了文档自动化处理领域的研究成果。"]
                            },
                            {
                                "title": "现实意义",
                                "paragraphs": ["该系统能够显著提高毕业论文排版效率。"]
                            }
                        ]
                    }
                ]
            },
            {
                "title": "系统设计与实现",
                "sections": [
                    {
                        "title": "系统架构",
                        "paragraphs": ["系统采用模块化设计，分为数据解析层、格式引擎层和输出生成层。"]
                    }
                ]
            }
        ],
        "references": [
            "重庆财经学院. 毕业论文（设计）指导手册[Z]. 2024.",
            "张三. 文档自动化处理技术研究[J]. 计算机应用, 2023, 43(5): 123-130.",
            "李四, 王五. 基于python-docx的Word文档生成方法[P]. 中国专利: CN202310123456.7, 2023."
        ],
        "appendices": [
            {
                "label": "A",
                "title": "系统核心代码",
                "content": "附录内容：系统核心代码清单。"
            }
        ]
    }

    # 生成文档
    formatter = ThesisFormatter()
    formatter.build(sample_thesis)
    output = formatter.save("重财论文示例_v1.1.1.docx")
    print(f"文档已生成：{output}")
    print("请在Word中打开后按 Ctrl+A 全选，再按 F9 更新所有域（目录、页码）。")
